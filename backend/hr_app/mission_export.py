"""Export des missions — PDF, Excel, Word."""
import io

from django.http import HttpResponse
from django.utils import timezone

try:
    from openpyxl import Workbook
except ImportError:
    Workbook = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from weasyprint import HTML

from .mission_service import STATUS_LABELS, mission_days_count


def _mission_row(m):
    emp = m.employee
    return [
        m.mission_number or str(m.id),
        emp.matricule,
        emp.full_name,
        emp.department.name if emp.department else '-',
        m.title,
        m.destination,
        str(m.start_date),
        str(m.end_date),
        mission_days_count(m),
        STATUS_LABELS.get(m.status, m.status),
        str(m.budget_allocated),
        str(m.daily_allowance),
    ]


def export_missions_excel(missions_qs):
    if not Workbook:
        raise RuntimeError('openpyxl non disponible')
    wb = Workbook()
    ws = wb.active
    ws.title = 'Missions'
    ws.append([
        'N°', 'Matricule', 'Employé', 'Département', 'Objet', 'Lieu',
        'Début', 'Fin', 'Jours', 'Statut', 'Budget', 'Indemnités/j',
    ])
    for m in missions_qs.select_related('employee', 'employee__department'):
        ws.append(_mission_row(m))
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f'EXPORT_MISSIONS_{timezone.now().strftime("%Y%m%d")}.xlsx'
    return buf.getvalue(), fname


def export_missions_pdf(missions_qs):
    rows_html = ''
    for m in missions_qs.select_related('employee', 'employee__department'):
        cols = _mission_row(m)
        cells = ''.join(f'<td>{c}</td>' for c in cols)
        rows_html += f'<tr>{cells}</tr>'
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <style>body{{font-family:Arial;margin:20px;font-size:10px}}
    table{{border-collapse:collapse;width:100%}}th,td{{border:1px solid #ccc;padding:5px}}
    th{{background:#1a5f9e;color:#fff}}</style></head>
    <body><h1>Liste des missions — OTOMIA RH</h1>
    <p>Généré le {timezone.now().strftime('%d/%m/%Y %H:%M')}</p>
    <table><thead><tr>
    <th>N°</th><th>Matricule</th><th>Employé</th><th>Dép.</th><th>Objet</th><th>Lieu</th>
    <th>Début</th><th>Fin</th><th>Jours</th><th>Statut</th><th>Budget</th><th>Ind./j</th>
    </tr></thead><tbody>{rows_html or '<tr><td colspan="12">Aucune mission</td></tr>'}</tbody></table>
    </body></html>"""
    fname = f'EXPORT_MISSIONS_{timezone.now().strftime("%Y%m%d")}.pdf'
    return HTML(string=html).write_pdf(), fname


def export_missions_word(missions_qs):
    if not DocxDocument:
        raise RuntimeError('python-docx non disponible')
    doc = DocxDocument()
    doc.add_heading('Liste des missions — OTOMIA RH', 0)
    doc.add_paragraph(f'Généré le {timezone.now().strftime("%d/%m/%Y %H:%M")}')
    for m in missions_qs.select_related('employee', 'employee__department'):
        doc.add_heading(f'{m.mission_number or m.id} — {m.title}', level=2)
        doc.add_paragraph(
            f'{m.employee.full_name} ({m.employee.matricule}) | {m.destination} | '
            f'{m.start_date} → {m.end_date} | {STATUS_LABELS.get(m.status, m.status)}'
        )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f'EXPORT_MISSIONS_{timezone.now().strftime("%Y%m%d")}.docx'
    return buf.getvalue(), fname


def mission_export_response(qs, fmt):
    fmt = (fmt or 'xlsx').lower()
    if fmt in ('xlsx', 'excel'):
        data, name = export_missions_excel(qs)
        ct = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif fmt in ('pdf',):
        data, name = export_missions_pdf(qs)
        ct = 'application/pdf'
    elif fmt in ('docx', 'word'):
        data, name = export_missions_word(qs)
        ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        raise ValueError('Format non supporté (pdf, xlsx, docx).')
    response = HttpResponse(data, content_type=ct)
    response['Content-Disposition'] = f'attachment; filename="{name}"'
    return response
