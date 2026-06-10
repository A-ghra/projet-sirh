import os
from datetime import datetime

from django.conf import settings
from django.template.loader import render_to_string
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from weasyprint import HTML

from .models import CompanySettings
from .payroll_service import payslip_filename
from .branding import DEVELOPER_SIGNATURE
from .payslip_builder import build_payslip_context, currency_symbol


def _ensure_dir(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)


def _company():
    return CompanySettings.get_settings()


def _legal_line(company):
    parts = [f'RCCM: {company.rccm}', f'ID. NAT: {company.id_nat}', f'CNSS: {company.cnss_number}']
    if company.tax_number:
        parts.append(f'Impôt: {company.tax_number}')
    if company.vat_number:
        parts.append(f'TVA: {company.vat_number}')
    return ' | '.join(parts)


def _payslip_context(payroll, user=None):
    return build_payslip_context(payroll, user=user)


def generate_payslip_pdf(payroll, user=None):
    ctx = _payslip_context(payroll, user=user)
    html_content = render_to_string('hr_app/payslip_template.html', ctx)
    filename = payslip_filename(payroll, 'pdf')
    pdf_path = os.path.join(settings.MEDIA_ROOT, 'payslips', filename)
    _ensure_dir(pdf_path)
    HTML(string=html_content).write_pdf(pdf_path)
    return f'payslips/{filename}'


def _fmt_amount(ctx, value):
    return f"{float(value):,.2f} {ctx['currency']}"


def generate_payslip_excel(payroll, user=None):
    ctx = _payslip_context(payroll, user=user)
    company = ctx['company']
    emp = ctx['employee']
    p = payroll
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Bulletin de Paie'
    bold = Font(bold=True)
    header_fill = PatternFill('solid', fgColor='1A5F9E')
    white = Font(bold=True, color='FFFFFF')

    def section(title):
        ws.append([])
        row = ws.max_row + 1
        ws.append([title])
        ws[f'A{row}'].font = bold

    ws.append([company.company_acronym])
    ws.append([company.company_slogan])
    ws.append([f"{company.bulletin_title} — {ctx['month_label'].upper()} {ctx['year_label']}"])
    ws.append([f"N° {ctx['bulletin_number']}"])
    ws.append([f"Devise : {ctx['currency_label']}"])
    ws.append([f"Généré le : {ctx['issued_at'].strftime('%d/%m/%Y %H:%M')}"])
    section('EMPLOYEUR')
    for row in [
        ['Raison sociale', company.company_name],
        ['Sigle', company.company_acronym],
        ['Adresse postale', company.postal_address or company.headquarters_address],
        ['Siège social', company.headquarters_address],
        ['Ville / Province / Pays', f"{company.city}, {company.province}, {company.country}"],
        ['Téléphone', company.phone_primary],
        ['Email', company.email],
        ['RCCM', company.rccm],
        ['ID. NAT', company.id_nat],
        ['N° Impôt', company.tax_number or '-'],
        ['N° CNSS', company.cnss_number],
    ]:
        ws.append(row)
    section('EMPLOYÉ')
    for row in [
        ['Matricule', emp.matricule],
        ['Nom', emp.nom or emp.full_name],
        ['Postnom', emp.postnom or '-'],
        ['Prénom', emp.prenom or '-'],
        ['Sexe', emp.get_gender_display()],
        ['Fonction', emp.position],
        ['Grade', emp.grade or '-'],
        ['Département', ctx['department']],
        ['Direction', ctx['direction']],
        ['Date embauche', str(emp.hire_date)],
        ['Ancienneté', f"{ctx['seniority_years']} ans"],
        ['Contrat', emp.contract_type],
        ['N° CNSS', emp.cnss_number or '-'],
        ['N° fiscal', emp.fiscal_number or '-'],
    ]:
        ws.append(row)
    section('PÉRIODE DE PAIE')
    for row in [
        ['Mois', f"{ctx['month_label']} {ctx['year_label']}"],
        ['Jours ouvrables', p.days_working],
        ['Jours travaillés', p.days_worked],
        ['Jours absence', p.days_absent],
        ['Jours congé', p.days_leave],
        ['Heures supplémentaires', float(p.overtime_hours)],
        ['Taux H. sup.', float(p.overtime_rate)],
    ]:
        ws.append(row)
    section('GAINS')
    ws.append(['Libellé', f'Montant ({ctx["currency"]})'])
    for label, val in ctx['gains']:
        ws.append([label, float(val)])
    ws.append(['SOUS-TOTAL GAINS (BRUT)', float(p.gross_salary)])
    section('RETENUES')
    ws.append(['Libellé', f'Montant ({ctx["currency"]})'])
    for label, val in ctx['retenues']:
        ws.append([label, float(val)])
    ws.append(['SOUS-TOTAL RETENUES', float(p.total_retenues)])
    section('RÉCAPITULATIF')
    for row in [
        ['Salaire de base', float(p.salary_base)],
        ['Total primes', float(p.total_primes)],
        ['Total indemnités', float(p.total_indemnites)],
        ['Salaire brut', float(p.gross_salary)],
        ['Salaire imposable', float(p.taxable_salary)],
        ['Montant CNSS', float(p.cnss_salarie)],
        ['Montant IPR', float(p.irpp)],
        ['Total retenues', float(p.total_retenues)],
        ['NET À PAYER', float(p.net_salary)],
    ]:
        ws.append(row)
    section('HISTORIQUE CONGÉS')
    ws.append(['Solde précédent', float(ctx['leave_previous'])])
    ws.append(['Congés pris', float(ctx['leave_taken'])])
    ws.append(['Solde actuel', float(ctx['leave_current'])])
    section('HISTORIQUE ABSENCES')
    ws.append(['Retards', ctx['absence_late']])
    ws.append(['Absences justifiées', ctx['absence_justified']])
    ws.append(['Absences non justifiées', ctx['absence_unjustified']])
    ws.append([])
    ws.append([f"Hash vérification : {ctx['verification_hash'][:16]}..."])
    ws.append([company.bulletin_footer])
    ws.append([f"Généré par {ctx['generated_by_name']} — {ctx['system_version']}"])
    ws.append([DEVELOPER_SIGNATURE])
    filename = payslip_filename(payroll, 'xlsx')
    file_path = os.path.join(settings.MEDIA_ROOT, 'payslips/excel', filename)
    _ensure_dir(file_path)
    wb.save(file_path)
    return f'payslips/excel/{filename}'


def generate_payslip_word(payroll, user=None):
    ctx = _payslip_context(payroll, user=user)
    company = ctx['company']
    emp = ctx['employee']
    p = payroll
    doc = Document()
    doc.add_heading(company.company_acronym, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(company.company_slogan).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(company.company_name).alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading(f"{company.bulletin_title} — {ctx['month_label']} {ctx['year_label']}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"N° {ctx['bulletin_number']} | Devise : {ctx['currency']} | Émis le {ctx['issued_at'].strftime('%d/%m/%Y')}")

    def add_section(title):
        doc.add_heading(title, level=2)

    add_section('Employeur')
    for line in [
        f"Adresse : {company.headquarters_address}",
        f"{company.city}, {company.province}, {company.country}",
        f"Tél : {company.phone_primary} | {company.email}",
        ctx['legal_line'],
    ]:
        doc.add_paragraph(line)

    add_section('Employé')
    for line in [
        f"Matricule : {emp.matricule} | Nom : {emp.nom} | Postnom : {emp.postnom} | Prénom : {emp.prenom}",
        f"Sexe : {emp.get_gender_display()} | Fonction : {emp.position} | Grade : {emp.grade or '-'}",
        f"Département : {ctx['department']} | Direction : {ctx['direction']}",
        f"Embauche : {emp.hire_date} | Ancienneté : {ctx['seniority_years']} ans | Contrat : {emp.contract_type}",
        f"CNSS : {emp.cnss_number or '-'} | Fiscal : {emp.fiscal_number or '-'}",
    ]:
        doc.add_paragraph(line)

    add_section('Période')
    doc.add_paragraph(
        f"Jours ouvrables : {p.days_working} | Travaillés : {p.days_worked} | "
        f"Absences : {p.days_absent} | Congés : {p.days_leave} | H. sup. : {p.overtime_hours}h"
    )

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Libellé'
    table.rows[0].cells[1].text = f'Montant ({ctx["currency"]})'
    doc.add_paragraph('— GAINS —')
    for label, val in ctx['gains']:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _fmt_amount(ctx, val)
    row = table.add_row().cells
    row[0].text = 'SOUS-TOTAL GAINS'
    row[1].text = _fmt_amount(ctx, p.gross_salary)
    doc.add_paragraph('— RETENUES —')
    for label, val in ctx['retenues']:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _fmt_amount(ctx, val)
    row = table.add_row().cells
    row[0].text = 'TOTAL RETENUES'
    row[1].text = _fmt_amount(ctx, p.total_retenues)

    net_p = doc.add_paragraph()
    net_p.add_run(f"\nNET À PAYER : {_fmt_amount(ctx, p.net_salary)}\n").bold = True
    net_p.runs[0].font.size = Pt(16)
    net_p.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x9E)

    add_section('Historique congés')
    doc.add_paragraph(f"Solde précédent : {ctx['leave_previous']} j | Pris : {ctx['leave_taken']} j | Actuel : {ctx['leave_current']} j")
    add_section('Historique absences')
    doc.add_paragraph(
        f"Retards : {ctx['absence_late']} | Justifiées : {ctx['absence_justified']} | "
        f"Non justifiées : {ctx['absence_unjustified']}"
    )
    doc.add_paragraph(f"\nVérification : {ctx['verification_hash'][:20]}...")
    doc.add_paragraph(f"Signature RH : {company.hr_manager_name} ({company.hr_department})")
    doc.add_paragraph(f"Direction : {company.director_name}")
    doc.add_paragraph(f"\n{company.bulletin_footer}")
    doc.add_paragraph(
        f"Document généré par OTOMIA RH — {ctx['issued_at'].strftime('%d/%m/%Y %H:%M')} — "
        f"{ctx['generated_by_name']} — {ctx['system_version']}"
    )
    p = doc.add_paragraph(DEVELOPER_SIGNATURE)
    p.runs[0].italic = True
    filename = payslip_filename(payroll, 'docx')
    file_path = os.path.join(settings.MEDIA_ROOT, 'payslips/word', filename)
    _ensure_dir(file_path)
    doc.save(file_path)
    return f'payslips/word/{filename}'


def global_payroll_filename(month_date, ext):
    months_fr = {
        1: 'JANVIER', 2: 'FEVRIER', 3: 'MARS', 4: 'AVRIL',
        5: 'MAI', 6: 'JUIN', 7: 'JUILLET', 8: 'AOUT',
        9: 'SEPTEMBRE', 10: 'OCTOBRE', 11: 'NOVEMBRE', 12: 'DECEMBRE',
    }
    mois = months_fr.get(month_date.month, 'MOIS')
    return f'ETAT_GLOBAL_PAIE_{mois}_{month_date.year}.{ext}'


def generate_payroll_mass_pdf(month_date, payrolls):
    company = _company()
    period = month_date.strftime('%B %Y').upper()
    rows = ''
    total_brut = total_net = total_cnss = total_irpp = 0
    for p in payrolls:
        curr = currency_symbol(p.currency)
        total_brut += float(p.gross_salary)
        total_net += float(p.net_salary)
        total_cnss += float(p.cnss_salarie)
        total_irpp += float(p.irpp)
        rows += f"""<tr>
            <td>{p.employee.matricule}</td><td>{p.employee.full_name}</td>
            <td>{p.gross_salary} {curr}</td><td>{p.cnss_salarie}</td><td>{p.irpp}</td>
            <td><strong>{p.net_salary} {curr}</strong></td><td>{p.get_status_display()}</td>
        </tr>"""
    html = f"""
    <html><head><style>
    body {{ font-family: Arial, sans-serif; margin: 30px; color: #2c3e50; }}
    .header {{ text-align: center; border-bottom: 3px solid #1a5f9e; padding-bottom: 12px; }}
    .header h1 {{ color: #1a5f9e; margin: 0; }}
    .meta {{ margin: 15px 0; font-size: 11px; color: #555; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 10px; }}
    th {{ background: #1a5f9e; color: white; padding: 8px; text-align: left; }}
    td {{ border: 1px solid #d5e6f7; padding: 7px; }}
    .totals {{ margin-top: 15px; background: #eaf4ff; padding: 12px; border-radius: 6px; }}
    .footer {{ margin-top: 25px; text-align: center; font-size: 9px; color: #999; }}
    </style></head><body>
    <div class="header">
        <h1>{company.company_acronym}</h1>
        <p>{company.company_slogan}</p>
        <h2>ÉTAT GLOBAL DE PAIE — {period}</h2>
        <p style="font-size:10px;color:#888;">Document récapitulatif collectif — Ne pas confondre avec un bulletin individuel</p>
    </div>
    <div class="meta">
        <strong>{company.company_name}</strong> — {company.headquarters_address}<br>
        {_legal_line(company)}
    </div>
    <table>
        <thead><tr>
            <th>Matricule</th><th>Employé</th><th>Brut</th>
            <th>CNSS</th><th>IRPP</th><th>Net</th><th>Statut</th>
        </tr></thead>
        <tbody>{rows}</tbody>
    </table>
    <div class="totals">
        <strong>Effectif :</strong> {len(payrolls)} employé(s) &nbsp;|&nbsp;
        <strong>Masse brute :</strong> {total_brut:,.2f} &nbsp;|&nbsp;
        <strong>CNSS :</strong> {total_cnss:,.2f} &nbsp;|&nbsp;
        <strong>IRPP :</strong> {total_irpp:,.2f} &nbsp;|&nbsp;
        <strong>Masse nette :</strong> {total_net:,.2f}
    </div>
    <div class="footer">{company.bulletin_footer}<br>Émis le {datetime.now().strftime('%d/%m/%Y %H:%M')} — {company.company_acronym}<br><em>{DEVELOPER_SIGNATURE}</em></div>
    </body></html>"""
    filename = global_payroll_filename(month_date, 'pdf')
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    _ensure_dir(file_path)
    HTML(string=html).write_pdf(file_path)
    return f'exports/{filename}'


def generate_payroll_mass_word(month_date, payrolls):
    company = _company()
    period = month_date.strftime('%B %Y').upper()
    doc = Document()
    doc.add_heading(company.company_acronym, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(company.company_slogan).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'ÉTAT GLOBAL DE PAIE — {period}', level=1)
    doc.add_paragraph('Document récapitulatif collectif — Ne pas confondre avec un bulletin individuel')
    doc.add_paragraph(f'{company.company_name} — {company.headquarters_address}')
    doc.add_paragraph(_legal_line(company))
    table = doc.add_table(rows=1, cols=7)
    headers = ['Matricule', 'Employé', 'Brut', 'CNSS', 'IRPP', 'Net', 'Statut']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    total_brut = total_net = 0
    for p in payrolls:
        row = table.add_row().cells
        row[0].text = p.employee.matricule
        row[1].text = p.employee.full_name
        row[2].text = str(p.gross_salary)
        row[3].text = str(p.cnss_salarie)
        row[4].text = str(p.irpp)
        row[5].text = str(p.net_salary)
        row[6].text = p.get_status_display()
        total_brut += float(p.gross_salary)
        total_net += float(p.net_salary)
    doc.add_paragraph(f'\nEffectif: {len(payrolls)} | Masse brute: {total_brut:,.2f} | Masse nette: {total_net:,.2f}')
    doc.add_paragraph(f'{company.bulletin_footer} — {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph(DEVELOPER_SIGNATURE)
    filename = global_payroll_filename(month_date, 'docx')
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    _ensure_dir(file_path)
    doc.save(file_path)
    return f'exports/{filename}'


def generate_payroll_excel(month_date, payrolls):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Paie {month_date.strftime('%m-%Y')}"
    company = _company()
    ws.append([company.company_acronym, company.company_name])
    ws.append([company.company_slogan])
    ws.append([f'ÉTAT GLOBAL DE PAIE — {month_date.strftime("%m/%Y")}'])
    ws.append(['Document récapitulatif collectif'])
    ws.append([_legal_line(company)])
    ws.append([])
    headers = [
        'Matricule', 'Nom complet', 'Devise', 'Brut', 'Primes', 'Indemnités',
        'Imposable', 'INPP', 'CNSS', 'IRPP', 'Retenues', 'Net', 'Statut'
    ]
    ws.append(headers)
    for p in payrolls:
        ws.append([
            p.employee.matricule, p.employee.full_name, p.currency,
            float(p.gross_salary), float(p.total_primes), float(p.total_indemnites),
            float(p.taxable_salary), float(p.inpp), float(p.cnss_salarie), float(p.irpp),
            float(p.total_retenues), float(p.net_salary), p.get_status_display(),
        ])
    filename = global_payroll_filename(month_date, 'xlsx')
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    _ensure_dir(file_path)
    wb.save(file_path)
    return f'exports/{filename}'


def generate_report_pdf(title, stats, author=None):
    company = _company()
    logo_html = ''
    from .company_utils import logo_file_uri
    logo_uri = logo_file_uri(company)
    if logo_uri:
        logo_html = f'<img src="{logo_uri}" style="max-height:55px;margin-bottom:8px;">'
    report_title = title or company.report_title
    author = author or company.report_author
    html = f"""
    <html><head><style>
    body {{ font-family: Arial, sans-serif; margin: 40px; color: #2c3e50; }}
    .header {{ text-align: center; border-bottom: 3px solid #1a5f9e; padding-bottom: 15px; }}
    .header h1 {{ color: #1a5f9e; margin: 0; }}
    .meta {{ font-size: 10px; color: #666; margin: 10px 0; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 30px; }}
    th, td {{ border: 1px solid #ddd; padding: 10px; }}
    th {{ background: #eaf4ff; }}
    .footer {{ margin-top: 30px; text-align: center; font-size: 9px; color: #999; }}
    </style></head><body>
    <div class="header">{logo_html}<h1>{company.company_acronym}</h1>
    <p>{company.report_header}</p><p>{company.report_subtitle}</p><h2>{report_title}</h2></div>
    <div class="meta">{_legal_line(company)}<br>{company.headquarters_address} | {company.phone_primary} | {company.email}<br>
    Auteur : {author} | Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}</div>
    <table><tr><th>Indicateur</th><th>Valeur</th></tr>
    """
    for key, value in stats.items():
        html += f"<tr><td>{key}</td><td>{value}</td></tr>"
    html += f"</table><div class='footer'>{company.report_footer}<br><em>{DEVELOPER_SIGNATURE}</em></div></body></html>"
    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports/pdf', filename)
    _ensure_dir(file_path)
    HTML(string=html).write_pdf(file_path)
    return f'reports/pdf/{filename}'


def generate_report_excel(title, stats):
    company = _company()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([company.company_acronym, company.company_name])
    ws.append([company.report_subtitle])
    ws.append([title or company.report_title])
    ws.append([_legal_line(company)])
    ws.append(['Indicateur', 'Valeur'])
    for key, value in stats.items():
        ws.append([key, value])
    ws.append([])
    ws.append([DEVELOPER_SIGNATURE])
    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports/excel', filename)
    _ensure_dir(file_path)
    wb.save(file_path)
    return f'reports/excel/{filename}'


def generate_report_word(title, stats):
    company = _company()
    doc = Document()
    doc.add_heading(company.company_acronym, 0)
    doc.add_paragraph(company.report_subtitle)
    doc.add_heading(title or company.report_title, level=1)
    doc.add_paragraph(_legal_line(company))
    doc.add_paragraph(company.report_footer)
    doc.add_paragraph(DEVELOPER_SIGNATURE)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Indicateur'
    table.rows[0].cells[1].text = 'Valeur'
    for key, value in stats.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)
    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports/word', filename)
    _ensure_dir(file_path)
    doc.save(file_path)
    return f'reports/word/{filename}'
