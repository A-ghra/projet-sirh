"""Export des contrats — PDF, Word, Excel avec archivage automatique."""
import io
import zipfile

from django.core.files.base import ContentFile
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font
except ImportError:
    Workbook = None

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    DocxDocument = None

from weasyprint import HTML

from .contract_builder import build_contract_context
from .contract_service import export_filename, log_contract_archive


def _fmt_money(val, currency='USD'):
    try:
        return f'{float(val):,.2f} {currency}'
    except (TypeError, ValueError):
        return str(val)


def _archive_export_file(contract, user, data, filename, fmt, bulk=False):
    """Enregistre le fichier exporté dans le journal d'archivage."""
    metadata = {'format': fmt, 'bulk': bulk, 'filename': filename}
    log = log_contract_archive(contract, 'EXPORT', user, note=filename, metadata=metadata)
    try:
        log.file_snapshot.save(filename, ContentFile(data), save=True)
    except Exception:
        pass
    return log


def _build_http_response(data, filename, content_type, export_type='individual'):
    response = HttpResponse(data, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response['X-Export-Type'] = export_type
    response['Cache-Control'] = 'no-store'
    return response


def export_contract_pdf(contract, user=None):
    ctx = build_contract_context(contract, user=user)
    html = render_to_string('hr_app/contract_template.html', ctx)
    pdf = HTML(string=html).write_pdf()
    return pdf, export_filename(contract, 'pdf')


def export_contract_excel(contract, user=None):
    if not Workbook:
        raise RuntimeError('openpyxl non disponible')
    ctx = build_contract_context(contract, user=user)
    company = ctx['company']
    emp = ctx['employee']
    c = contract
    wb = Workbook()
    ws = wb.active
    ws.title = 'Contrat'
    bold = Font(bold=True)

    def section(title):
        ws.append([])
        row = ws.max_row + 1
        ws.append([title])
        ws[f'A{row}'].font = bold

    ws.append([company.company_acronym])
    ws.append([company.company_name])
    ws.append([f'CONTRAT DE TRAVAIL — {c.contract_number}'])
    ws.append([f'Généré le : {ctx["issued_at"].strftime("%d/%m/%Y %H:%M")}'])
    section('ENTREPRISE')
    for row in [
        ['Raison sociale', company.company_name],
        ['RCCM', company.rccm],
        ['ID. NAT', company.id_nat],
        ['CNSS', company.cnss_number],
        ['Adresse', company.headquarters_address],
        ['Téléphone', company.phone_primary],
        ['Email', company.email],
    ]:
        ws.append(row)
    section('EMPLOYÉ')
    for row in [
        ['Matricule', emp.matricule],
        ['Nom complet', emp.full_name],
        ['Sexe', emp.get_gender_display()],
        ['Nationalité', emp.nationality],
        ['Adresse', emp.address or '—'],
        ['Département', ctx['department']],
        ['Poste', c.position_title or emp.position],
        ['Responsable', ctx['manager_name']],
    ]:
        ws.append(row)
    section('CONTRAT')
    for row in [
        ['Type', c.contract_type],
        ['Date début', str(c.start_date)],
        ['Date fin', str(c.end_date or '—')],
        ['Durée', ctx['duration_label']],
        ['Période essai (fin)', str(c.probation_end_date or '—')],
        ['Jours/semaine', c.work_days_per_week or 5],
        ['Horaires', c.work_schedule or '—'],
        ['Lieu affectation', c.assignment_location or '—'],
    ]:
        ws.append(row)
    section('RÉMUNÉRATION')
    for label, val in ctx['remuneration_lines']:
        ws.append([label, _fmt_money(val, c.currency)])
    for label, val in ctx['benefits_lines']:
        ws.append([f'{label} (avantage)', _fmt_money(val, c.currency)])
    section('CLAUSES')
    for label, val in [
        ('Obligations employé', c.employee_obligations),
        ('Obligations employeur', c.employer_obligations),
        ('Résiliation', c.termination_conditions),
        ('Confidentialité', c.confidentiality_clause),
        ('Références légales', ctx['labor_code_refs']),
    ]:
        if val:
            ws.append([label, val])
    section('SIGNATURES')
    ws.append(['Employé', '✅' if ctx['employee_signed'] else 'En attente'])
    ws.append(['RH', c.hr_signatory_name or 'En attente'])
    ws.append(['Direction', c.direction_signatory_name or 'En attente'])
    ws.append(['Hash vérification', ctx['verification_hash'][:20] + '...'])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue(), export_filename(contract, 'xlsx')


def export_contract_word(contract, user=None):
    if not DocxDocument:
        raise RuntimeError('python-docx non disponible')
    ctx = build_contract_context(contract, user=user)
    company = ctx['company']
    emp = ctx['employee']
    c = contract
    doc = DocxDocument()
    title = doc.add_heading(f'{company.company_acronym} — Contrat de travail', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(company.company_name)
    doc.add_paragraph(f'RCCM: {company.rccm} | ID. NAT: {company.id_nat} | CNSS: {company.cnss_number}')
    doc.add_paragraph(f'{company.headquarters_address} — {company.phone_primary} — {company.email}')
    doc.add_heading(f'Contrat N° {c.contract_number} — {c.contract_type}', level=1)

    doc.add_heading('Employé', level=2)
    for label, val in [
        ('Matricule', emp.matricule), ('Nom', emp.full_name),
        ('Sexe', emp.get_gender_display()), ('Nationalité', emp.nationality),
        ('Adresse', emp.address or '—'), ('Département', ctx['department']),
        ('Poste', c.position_title or emp.position), ('Responsable', ctx['manager_name']),
    ]:
        doc.add_paragraph(f'{label} : {val}')

    doc.add_heading('Informations contractuelles', level=2)
    for label, val in [
        ('Type', c.contract_type), ('Début', str(c.start_date)),
        ('Fin', str(c.end_date or '—')), ('Durée', ctx['duration_label']),
        ('Essai (fin)', str(c.probation_end_date or '—')),
        ('Jours/semaine', c.work_days_per_week or 5),
        ('Horaires', c.work_schedule or '—'), ('Lieu', c.assignment_location or '—'),
    ]:
        doc.add_paragraph(f'{label} : {val}')

    doc.add_heading('Rémunération', level=2)
    for label, val in ctx['remuneration_lines']:
        doc.add_paragraph(f'{label} : {_fmt_money(val, c.currency)}')
    for label, val in ctx['benefits_lines']:
        doc.add_paragraph(f'{label} : {_fmt_money(val, c.currency)}')

    doc.add_heading('Clauses', level=2)
    for label, val in [
        ('Obligations employé', c.employee_obligations),
        ('Obligations employeur', c.employer_obligations),
        ('Résiliation', c.termination_conditions),
        ('Confidentialité', c.confidentiality_clause),
        ('Références légales', ctx['labor_code_refs']),
    ]:
        if val:
            doc.add_heading(label, level=3)
            doc.add_paragraph(val)

    doc.add_heading('Signatures', level=2)
    doc.add_paragraph(f"Employé : {'Signé' if ctx['employee_signed'] else 'En attente'} — {emp.full_name}")
    doc.add_paragraph(f"RH : {c.hr_signatory_name or 'En attente'}")
    doc.add_paragraph(f"Direction : {c.direction_signatory_name or 'En attente'}")
    doc.add_paragraph(f"Vérification : {ctx['verification_hash'][:20]}...")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), export_filename(contract, 'docx')


def contract_export_response(contract, fmt, user=None, export_type='individual'):
    from .models import ContractDownloadLog

    fmt = (fmt or 'pdf').lower()
    if fmt in ('xlsx', 'excel'):
        data, name = export_contract_excel(contract, user=user)
        content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif fmt in ('docx', 'word'):
        data, name = export_contract_word(contract, user=user)
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        data, name = export_contract_pdf(contract, user=user)
        content_type = 'application/pdf'

    if user and user.is_authenticated:
        ContractDownloadLog.objects.create(contract=contract, user=user, format=fmt)
        _archive_export_file(contract, user, data, name, fmt, bulk=False)

    return _build_http_response(data, name, content_type, export_type=export_type)


def export_contracts_bulk_pdf(contracts_qs, user=None):
    sections = []
    for i, c in enumerate(contracts_qs.select_related('employee', 'employee__department', 'employee__manager')):
        ctx = build_contract_context(c, user=user)
        html_part = render_to_string('hr_app/contract_template.html', ctx)
        body_match = html_part.split('<body>')[-1].split('</body>')[0] if '<body>' in html_part else html_part
        brk = ' style="page-break-before:always"' if i else ''
        sections.append(f'<div{brk}>{body_match}</div>')
    full_html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <style>@page{{size:A4;margin:12mm}}body{{font-family:Arial}}</style></head>
    <body>{''.join(sections) if sections else '<p>Aucun contrat.</p>'}</body></html>"""
    pdf = HTML(string=full_html).write_pdf()
    fname = f'EXPORT_CONTRATS_{timezone.now().strftime("%Y%m%d")}.pdf'
    return pdf, fname


def export_contracts_bulk_zip_docx(contracts_qs, user=None):
    buf = io.BytesIO()
    used_names = set()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for c in contracts_qs.select_related('employee'):
            try:
                data, name = export_contract_word(c, user=user)
                if name in used_names:
                    name = f'{c.contract_number}_{name}'
                used_names.add(name)
                zf.writestr(name, data)
            except Exception:
                continue
    buf.seek(0)
    fname = f'EXPORT_CONTRATS_{timezone.now().strftime("%Y%m%d")}.zip'
    return buf.getvalue(), fname


def export_contracts_bulk(contracts_qs, fmt='xlsx'):
    if not Workbook:
        raise RuntimeError('openpyxl non disponible')
    wb = Workbook()
    ws = wb.active
    ws.title = 'Contrats'
    ws.append([
        'N° contrat', 'Matricule', 'Employé', 'Département', 'Type',
        'Début', 'Fin', 'Statut', 'Statut métier', 'Salaire', 'Devise',
    ])
    from .contract_service import compute_lifecycle_status, LIFECYCLE_LABELS
    for c in contracts_qs.select_related('employee', 'employee__department'):
        ws.append([
            c.contract_number, c.employee.matricule, c.employee.full_name,
            c.employee.department.name if c.employee.department else '-',
            c.contract_type, str(c.start_date), str(c.end_date or ''),
            c.get_status_display(),
            LIFECYCLE_LABELS.get(compute_lifecycle_status(c), ''),
            float(c.salary_base), c.currency,
        ])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f'EXPORT_CONTRATS_{timezone.now().strftime("%Y%m%d")}.xlsx'
    return buf.getvalue(), fname


def export_global_response(qs, fmt, user=None):
    fmt = (fmt or 'xlsx').lower()
    if fmt in ('xlsx', 'excel'):
        data, fname = export_contracts_bulk(qs, fmt)
        content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif fmt in ('pdf',):
        data, fname = export_contracts_bulk_pdf(qs, user=user)
        content_type = 'application/pdf'
    elif fmt in ('docx', 'word'):
        data, fname = export_contracts_bulk_zip_docx(qs, user=user)
        content_type = 'application/zip'
    else:
        raise ValueError('Format non supporté (pdf, xlsx, docx).')

    if user and user.is_authenticated:
        for c in qs[:50]:
            _archive_export_file(c, user, data, fname, fmt, bulk=True)

    return _build_http_response(data, fname, content_type, export_type='global')
